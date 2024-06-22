import re
import os
import psutil
import pyperclip
import multiprocessing

from functools import partial
from PyPDF2 import PdfReader, PdfWriter
from docx2pdf import convert
from docx import Document
from datetime import datetime
from openpyxl import load_workbook
from openpyxl.styles import Font


from selenium_script import FirefoxBrowser, copy_cookies_db


class ResumeHelper:
    def __init__(self):
        self.process_to_kill = [
            "WINWORD.EXE",
            "FoxitPDFEditor.exe",
            # "Notepad.exe",
            "Acrobat.exe",
            "EXCEL.EXE",
        ]
        # Get the current file directory
        current_file_dir = os.path.dirname(os.path.realpath(__file__))

        # Go one directory up
        parent_dir = os.path.dirname(current_file_dir)

        self.template_path = os.path.join(parent_dir, "templates")

        self.cover_letter_output = (
            f"{parent_dir}\\output_files\\Ahmed_Qureshi_Cover_Letter.docx"
        )
        self.cover_letter_txt = f"{parent_dir}\\output_files\\cover_letter.txt"

        self.resume_output = f"{parent_dir}\\output_files\\Ahmed_Qureshi_Resume.docx"
        self.cover_letter_template = (
            f"{self.template_path}\\Ahmed_Qureshi_Cover_Letter_Template.docx"
        )

        # self.cover_letter_template = f"{self.template_path}\\cover_letter_template.txt"
        self.resume_summary_template = (
            f"{self.template_path}\\resume_summary_template.txt"
        )

        self.resume_pdf = (
            r"C:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed Qureshi Resume.pdf"
        )
        self.cover_letter_pdf = (
            r"C:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed Qureshi Cover Letter.pdf"
        )

        # self.cover_resume_proj = (
        #     r"C:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed_Qureshi_Resume_ProjPortfolio.pdf",
        # )
        self.resume_cover_merge = (
            r"c:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed_Qureshi_Cover_Resume.pdf"
        )

        self.job_tracker_xls = (
            f"{parent_dir}\\output_files\\job_application_tracker_company_list.xlsx"
        )

    def default_resume_path(self):
        self.resume_template = (
            f"{self.template_path}\\Ahmed_Qureshi_Resume_Template.docx"
        )
        self.resume_txt = f"{self.template_path}\\resume.txt"
        print(f"\nThis is the path {self.resume_template}\n")

    def close_apps(self):
        for process in (
            process
            for process in psutil.process_iter()
            if process.name() in self.process_to_kill
        ):
            process.kill()

    def add_row_to_excel(self, file_path, data_list):
        # Load the workbook and select the active worksheet
        wb = load_workbook(filename=file_path)
        ws = wb.active

        # Append the data to the end of the worksheet
        ws.append(data_list)

        # Get the last row number
        last_row = ws.max_row

        # Change the font size of the last row to 12
        for cell in ws[last_row]:
            cell.font = Font(size=12)

        # Save the workbook
        wb.save(file_path)

    def format_data(self, data):
        # Remove special characters
        data = re.sub(r"\W+", " ", data)

        # Replace multiple spaces with a single space
        data = re.sub(r"\s+", " ", data)

        # Remove line breaks
        data = data.replace("\n", " ")

        return data

    def replace_lines_breaks(self, data):
        # Remove extra spaces
        text = re.sub(" +", " ", data)

        # Replace multiple line breaks with a single one
        text = re.sub("\r\n\r\n", "\n\n", text)
        text = re.sub("\\n+", "\\n\\n", text)

        return text

    def format_text_single_break(self, text):
        # Remove special characters but keep commas
        text = re.sub(r"[^\w\s\n,]", "", text)

        # Replace multiple spaces with a single space
        text = re.sub(r" +", " ", text)

        # Replace multiple line breaks with a single line break
        text = re.sub("\r\n\r\n", "\n", text)
        text = re.sub("\r\n", "\n", text)
        text = re.sub(r"\n+", "\n", text)

        return text

    def copy_keyword_job_resume(self, filename, job_description):
        with open(filename, "r") as file:
            data = file.read()

        data += job_description

        with open(self.resume_txt, "r") as file:
            data += file.read()

        data = self.format_data(data)

        return data
        # Copy to clipboard

        # pyperclip.copy(data)

    def replace_string_in_word_table(
        self,
        search_text,
        replacement_text,
        file,
        save_location,
        itables=0,
        irow=3,
        icell=0,
    ):
        doc = Document(file)

        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             for run in cell.paragraphs:
        #                 print(run.text)

        cell = doc.tables[itables].rows[irow].cells[icell]
        for i in range(len(cell.paragraphs)):
            for run in cell.paragraphs[i].runs:
                if search_text.lower() in run.text.lower():
                    run.text = run.text.replace(search_text, replacement_text.strip())

        doc.save(save_location)

    def generate_pdf(self, docx_file, pdf_output):
        self.close_apps()
        convert(docx_file, pdf_output)

    def merge_pdfs(self, pdf1, pdf2, save_location):
        # Create a PDF file writer object
        writer = PdfWriter()

        # Add the pages from the first file
        for page in PdfReader(pdf1).pages:
            writer.add_page(page)

        # Add the pages from the second file
        for page in PdfReader(pdf2).pages:
            writer.add_page(page)

        # Write the result to the save_location file
        with open(save_location, "wb") as save_location_file:
            writer.write(save_location_file)

        os.startfile(save_location)

    def launch_file(self, filename):
        os.startfile(filename)

    def get_formated_date(self):
        # Get today's date
        now = datetime.now()

        # Define function to get the ordinal suffix
        def get_ordinal(n):
            if 10 <= n <= 20:
                return "th"
            else:
                return {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")

        # Format the date
        day = now.day
        formatted_date = f"{day}{get_ordinal(day)} {now.strftime('%B, %Y')}"

        return formatted_date

    def run_in_multiprocessing(self, func, *args):
        copy_cookies_db()
        pool = multiprocessing.Pool(processes=2)

        for arg_tuple in args:
            partial_func = partial(func, *arg_tuple)
            pool.apply_async(partial_func)

        pool.close()
        pool.join()

    def run_in_multiprocessing_without_partial(self, func, *args):

        process = multiprocessing.Process(target=func, args=args)
        process.start()
        process.join()

    def get_programmer_title(self):
        title_dict = {
            1: "Web Developer",
            2: "Front-End Developer",
            3: "Back-End Developer",
            4: "Full Stack Developer",
            5: "Software Developer",
            6: "Python Developer",
            7: "Python Programmer",
            8: "JavaScript Developer",
            9: "React Developer",
            10: "Flutter Developer",
            11: "Programmer",
            12: "It Help Desk Analyst",
            13: "Administration Clerk",
            # 11: "Mobile App Developer",
            # 12: "DevOps Engineer",
            # 13: "System Administrator",
            # 14: "Database Administrator",
        }
        resume_type_mapping = {
            12: ("It_Support", "resume_it_support"),
            13: ("Admin_Clerk", "resume_admin_clerk"),
        }

        while True:
            print("Please choose a number from the following options:")
            print("\n".join(f"{num}: {title}" for num, title in title_dict.items()))
            choice = input("\nEnter your choice: ")

            if choice.isdigit() and int(choice) in title_dict:
                number = int(choice)
                if number in resume_type_mapping:
                    template, txt_file = resume_type_mapping[number]
                    self.resume_template = f"{self.template_path}\\Ahmed_Qureshi_{template}_Resume_Template.docx"
                    self.resume_txt = f"{self.template_path}\\{txt_file}.txt"

                return title_dict[number]

            if not choice.isdigit():
                resume_type = input(
                    "\nSelect Resume Type\n1: Programmer\n2: IT Support\n3: Administration Clerk"
                )
                if resume_type in ("1", "2", "3"):
                    if resume_type == "2":
                        self.resume_template = f"{self.template_path}\\Ahmed_Qureshi_It_Support_Resume_Template.docx"
                        self.resume_txt = f"{self.template_path}\\resume_it_support.txt"
                    elif resume_type == "3":
                        self.resume_template = f"{self.template_path}\\Ahmed_Qureshi_Admin_Clerk_Resume_Template.docx"
                        self.resume_txt = (
                            f"{self.template_path}\\resume_admin_clerk.txt"
                        )
                    return choice

            print("\nInvalid input. Please try again.\n")

    def replace_string_word(
        self,
        search_text,
        replacement_text: str,
        file,
        save_location,
        isbold: bool = False,
    ):
        doc = Document(file)
        for p in doc.paragraphs:
            if search_text in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if search_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(
                            search_text, replacement_text.strip()
                        )
                        if isbold:
                            inline[i].bold = True

        doc.save(save_location)

    def copy_word_to_txt(self, word_file_path, txt_file_path):
        # Read the content from the Word file
        doc = Document(word_file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        word_content = "\n".join(full_text)

        # Write the content to the txt file (this will overwrite existing content)
        with open(txt_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(word_content)

        self.launch_file(txt_file_path)

    def run(self):
        run = True
        while run:
            self.default_resume_path()

            designaion = self.get_programmer_title()

            date_now = self.get_formated_date()

            # Resume Workings
            self.replace_string_in_word_table(
                search_text="date",
                replacement_text=date_now,
                file=self.resume_template,
                save_location=self.resume_output,
                irow=4,
                icell=1,
            )

            self.replace_string_word(
                search_text="designation",
                replacement_text=designaion,
                file=self.resume_output,
                save_location=self.resume_output,
            )

            input("\n\nPlease Copy Contents of Job Description and Press Enter ")

            job_description = self.format_data(pyperclip.paste())

            self.replace_string_word(
                search_text="job_description",
                replacement_text=job_description,
                file=self.resume_output,
                save_location=self.resume_output,
            )

            search = self.copy_keyword_job_resume(
                self.resume_summary_template, job_description
            )

            self.run_in_multiprocessing(
                FirefoxBrowser,
                (search, "chatgpt"),
                (search, "gemini"),
            )

            input("\n\n✅ Please Copy The Summary and Press Enter ")

            self.replace_string_word(
                search_text="summary_placeholder",
                replacement_text=pyperclip.paste(),
                file=self.resume_output,
                save_location=self.resume_output,
            )

            self.launch_file(self.resume_output)

            input("\n\nPress Enter Generate Resume PDF")

            self.run_in_multiprocessing_without_partial(
                self.generate_pdf, self.resume_output, self.resume_pdf
            )
            # Cover Leter
            self.replace_string_word(
                search_text="designation",
                replacement_text=designaion,
                file=self.cover_letter_template,
                save_location=self.cover_letter_output,
            )

            self.replace_string_word(
                search_text="date",
                replacement_text=date_now,
                file=self.cover_letter_output,
                save_location=self.cover_letter_output,
            )

            input("\n\nPress Enter, Cover Body On The Word Template ")

            self.replace_string_word(
                search_text="body",
                replacement_text=self.replace_lines_breaks(pyperclip.paste()),
                file=self.cover_letter_output,
                save_location=self.cover_letter_output,
            )

            input("\n\nPress Enter To Copy the Company Name On The Word Template ")
            company_name = self.format_text_single_break(pyperclip.paste())
            self.replace_string_word(
                search_text="company_name",
                replacement_text=company_name,
                file=self.cover_letter_output,
                save_location=self.cover_letter_output,
                isbold=True,
            )
            contact_person = self.format_data(
                input("\n\n⚠️  Enter a Contact Person Manually ")
            )
            if contact_person:
                self.replace_string_word(
                    search_text="Hiring Manager",
                    replacement_text=contact_person,
                    file=self.cover_letter_output,
                    save_location=self.cover_letter_output,
                    isbold=False,
                )
            else:
                contact_person = "Hiring Manager"

            self.launch_file(self.cover_letter_output)

            input("\n\nPress Enter To Generate Cover Letter ")

            self.copy_word_to_txt(self.cover_letter_output, self.cover_letter_txt)

            self.generate_pdf(self.cover_letter_output, self.cover_letter_pdf)

            self.merge_pdfs(
                self.cover_letter_pdf, self.resume_pdf, self.resume_cover_merge
            )

            self.add_row_to_excel(
                self.job_tracker_xls,
                [
                    company_name,
                    designaion,
                    contact_person,
                    "",
                    "Remote",
                    "",
                    "",
                    date_now,
                ],
            )

            # self.launch_file(self.job_tracker_xls)

            continue_run = input(
                """
            \n\nAll Operations are Completed\n
            ---------------------------------
            |                               |
            |   Script execution complete   |
            |                               |
            ---------------------------------
            \n\nPress Enter to Start Again...
            """
            )
            if continue_run == "1":
                run = False


if __name__ == "__main__":
    helper = ResumeHelper()
    helper.run()
