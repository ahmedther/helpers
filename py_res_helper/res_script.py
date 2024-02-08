import re
import os
import psutil
import pyperclip

from PyPDF2 import PdfReader, PdfWriter
from docx2pdf import convert
from docx import Document
from datetime import datetime


class ResumeHelper:
    def __init__(self):
        self.process_to_kill = [
            "WINWORD.EXE",
            "FoxitPDFEditor.exe",
            # "Notepad.exe",
            "Acrobat.exe",
        ]
        self.template_path = f"{os.getcwd()}\\templates"
        self.cover_letter_output = (
            f"{os.getcwd()}\\output_files\\Ahmed_Qureshi_Cover_Letter.docx"
        )
        self.resume_output = f"{os.getcwd()}\\output_files\\Ahmed_Qureshi_Resume.docx"
        self.cover_letter_template = (
            f"{self.template_path}\Ahmed_Qureshi_Cover_Letter_Template.docx"
        )
        self.resume_template = (
            f"{self.template_path}\Ahmed_Qureshi_Resume_Template.docx"
        )
        # self.cover_letter_template = f"{self.template_path}\\cover_letter_template.txt"
        self.resume_summary_template = (
            f"{self.template_path}\\resume_summary_template.txt"
        )
        self.resume_txt = f"{self.template_path}\\resume.txt"

        self.resume_pdf = (
            r"C:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed Qureshi Resume.pdf"
        )
        self.cover_letter_pdf = (
            r"C:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed Qureshi Cover Letter.pdf",
        )
        # self.cover_resume_proj = (
        #     r"C:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed_Qureshi_Resume_ProjPortfolio.pdf",
        # )
        self.resume_cover_merge = (
            r"c:\Users\AHMED\Desktop\AHMED\Resume\pdf\Ahmed_Qureshi_Cover_Resume.pdf",
        )

    def close_apps(self):
        for process in (
            process
            for process in psutil.process_iter()
            if process.name() in self.process_to_kill
        ):
            process.kill()

    def format_data(self, data):
        # Remove special characters
        data = re.sub(r"\W+", " ", data)

        # Replace multiple spaces with a single space
        data = re.sub(r"\s+", " ", data)

        # Remove line breaks
        data = data.replace("\n", " ")

        return data

    def copy_keyword_job_resume(self, filename, job_description):
        with open(filename, "r") as file:
            data = file.read()

        data += job_description

        with open(self.resume_txt, "r") as file:
            data += file.read()

        self.format_data(data)
        # Copy to clipboard

        pyperclip.copy(data)

    def replace_string_in_word_table(
        self,
        search_text,
        replacement_text,
        file,
        save_location,
        itables=0,
        irow=3,
        icell=0,
    ):
        doc = Document(file)
        # Remove extra spaces
        text = re.sub(" +", " ", replacement_text)

        # Replace multiple line breaks with a single one
        text = re.sub("\r\n\r\n", "\n\n", text)
        text = re.sub("\\n+", "\\n\\n", text)

        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             for run in cell.paragraphs:
        #                 print(run.text)

        cell = doc.tables[itables].rows[irow].cells[icell]
        for i in range(len(cell.paragraphs)):
            for run in cell.paragraphs[i].runs:
                print(run.text)
                if search_text.lower() in run.text.lower():
                    run.text = run.text.replace(search_text, text.strip())

        doc.save(save_location)

    def generate_pdf(self, docx_file, pdf_output):

        convert(docx_file, pdf_output)

    def merge_pdfs(self, pdf1, pdf2, save_location):
        # Create a PDF file writer object
        writer = PdfWriter()

        # Add the pages from the first file
        for page in PdfReader(pdf1).pages:
            writer.add_page(page)

        # Add the pages from the second file
        for page in PdfReader(pdf2).pages:
            writer.add_page(page)

        # Write the result to the save_location file
        with open(save_location, "wb") as save_location_file:
            writer.write(save_location_file)

        os.startfile(save_location)

    def launch_file(self, filename):
        os.startfile(filename)

    def get_formated_date(self):
        # Get today's date
        now = datetime.now()

        # Define function to get the ordinal suffix
        def get_ordinal(n):
            if 10 <= n <= 20:
                return "th"
            else:
                return {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")

        # Format the date
        day = now.day
        formatted_date = f"{day}{get_ordinal(day)} {now.strftime('%B, %Y')}"

        return formatted_date

    def get_programmer_title(self):
        title_dict = {
            1: "Web Developer",
            2: "Front-End Developer",
            3: "Back-End Developer",
            4: "Full Stack Developer",
            5: "Software Developer",
            6: "Python Developer",
            7: "Python Programmer",
            8: "JavaScript Developer",
            9: "React Developer",
            10: "Flutter Developer",
            11: "Programmer",
            # 11: "Mobile App Developer",
            # 12: "DevOps Engineer",
            # 13: "System Administrator",
            # 14: "Database Administrator",
        }

        print("Please choose a number from the following options:")
        for number, title in title_dict.items():
            print(f"{number}: {title}")

        number = int(input("\nEnter your choice: "))
        return title_dict.get(
            number, "Invalid number. Please enter a number between 1 and 12."
        )

    def replace_string_word(
        self, search_text, replacement_text: str, file, save_location
    ):
        doc = Document(file)

        for p in doc.paragraphs:
            if search_text in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if search_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(
                            search_text, replacement_text.strip()
                        )

        doc.save(save_location)

    def run(self):
        run = True
        open_notepad = input(
            "\n\nPress 1 to Open the Notepad Template. Otherwise Press Enter To Continue... "
        )
        if open_notepad == "1":
            self.launch_file(self.notepad_template)

        # input("\n\nPress Enter Process and Copy Data From the Notepad Template ")

        while run:
            self.close_apps()

            # Resume Workings
            self.replace_string_in_word_table(
                search_text="date",
                replacement_text=self.get_formated_date(),
                file=self.resume_template,
                save_location=self.resume_output,
                irow=5,
                icell=1,
            )

            designaion = self.get_programmer_title()

            self.replace_string_word(
                search_text="designation",
                replacement_text=designaion,
                file=self.resume_output,
                save_location=self.resume_output,
            )

            input("\n\nPlease Copy Contents of Job Description and Press Enter ")

            job_description = self.format_data(pyperclip.paste())

            self.copy_keyword_job_resume(self.resume_summary_template, job_description)

            input("\n\n✅ Copied!!! Please Copy The Summary and Press Enter ")

            self.replace_string_word(
                search_text="summaryplaceholder",
                replacement_text=pyperclip.paste(),
                file=self.resume_output,
                save_location=self.resume_output,
            )

            self.launch_file(self.resume_output)

            input("\n\nPress Enter Generate Resume PDF")

            self.generate_pdf(self.resume_output, self.resume_pdf)

            # Cover Leter
            self.replace_string_word(
                search_text="date",
                replacement_text=self.get_formated_date(),
                file=self.cover_letter_template,
                save_location=self.cover_letter_output,
            )

            input("\n\nPress Enter To Copy the Company Name On The Word Template ")

            self.replace_string_word(
                search_text="company",
                replacement_text=pyperclip.paste(),
                file=self.cover_letter_output,
                save_location=self.cover_letter_output,
            )

            input("\n\nPress Enter To Copy the Body On The Word Template ")

            self.replace_string_word(
                search_text="body",
                replacement_text=pyperclip.paste(),
                file=self.cover_letter_output,
                save_location=self.cover_letter_output,
            )

            input("\n\nPress Enter To Generate and Merge PDFs ")

            self.merge_pdfs(
                self.resume_pdf, self.cover_letter_pdf, self.resume_cover_merge
            )

            continue_run = input(
                """
\n\nAll Operations are Completed\n
---------------------------------
|                               |
|   Script execution complete   |
|                               |
---------------------------------
\n\nPress Enter to Start Again...
"""
            )
            if continue_run == "1":
                run = False


if __name__ == "__main__":
    helper = ResumeHelper()
    helper.run()
